"""Generate architecture description document for the edu_system project."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.expanduser("~/Desktop/Описание программы.docx")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Inches(1.18)   # 3 cm
section.right_margin  = Inches(0.79)   # 2 cm
section.top_margin    = Inches(0.79)
section.bottom_margin = Inches(0.79)

# ── Styles ────────────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)

h1_style = doc.styles["Heading 1"]
h1_style.font.name = "Arial"
h1_style.font.size = Pt(16)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0x26, 0x30, 0x45)  # vgtu-dark

h2_style = doc.styles["Heading 2"]
h2_style.font.name = "Arial"
h2_style.font.size = Pt(13)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(0x26, 0x30, 0x45)

h3_style = doc.styles["Heading 3"]
h3_style.font.name = "Arial"
h3_style.font.size = Pt(11)
h3_style.font.bold = True
h3_style.font.color.rgb = RGBColor(0x3a, 0x4a, 0x65)


def h1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def code_block(lines):
    """Render a small code snippet with Courier New."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    doc.add_paragraph()


def divider():
    doc.add_paragraph("─" * 80)


# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tp.add_run("ВОРОНЕЖСКИЙ ГОСУДАРСТВЕННЫЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ")
run.font.name = "Arial"
run.font.size = Pt(12)
run.font.bold = True
run.font.color.rgb = RGBColor(0x26, 0x30, 0x45)

doc.add_paragraph()

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = tp2.add_run("Информационная система управления\nобразовательными программами")
run2.font.name = "Arial"
run2.font.size = Pt(20)
run2.font.bold = True
run2.font.color.rgb = RGBColor(0x26, 0x30, 0x45)

doc.add_paragraph()

tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = tp3.add_run("Техническое описание архитектуры, стека и ключевых решений")
run3.font.name = "Arial"
run3.font.size = Pt(13)
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

tp4 = doc.add_paragraph()
tp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = tp4.add_run("Модуль: заведующий кафедрой\nВерсия 1.0   |   2026")
run4.font.name = "Arial"
run4.font.size = Pt(11)
run4.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. НАЗНАЧЕНИЕ СИСТЕМЫ
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Назначение системы")

para(
    "Система предназначена для автоматизации работы заведующего кафедрой ВГТУ "
    "в части управления образовательными программами и контроля отчётной документации "
    "по дисциплинам. Ключевые бизнес-процессы:"
)
bullet("Просмотр образовательных программ кафедры с детализацией по дисциплинам.")
bullet("Ведение чеклистов готовности отчётной документации по каждой дисциплине.")
bullet("Загрузка файлов отчётов (PDF / DOC / DOCX) и привязка их к разделам чеклиста.")
bullet("Маршрут утверждения: преподаватель подаёт → заведующий утверждает. "
       "Заведующий не может утвердить собственный отчёт — защита от конфликта интересов.")
bullet("Обратная связь по каждому файлу отчёта.")
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  2. ТЕХНОЛОГИЧЕСКИЙ СТЕК
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Технологический стек и мотивация выбора")

h2("2.1 Backend — FastAPI")
para(
    "FastAPI выбран как основной web-фреймворк. Причины:"
)
bullet(
    "Нативная поддержка async/await — критично, так как БД-драйвер asyncpg "
    "работает только в асинхронном режиме. Синхронный Django или Flask потребовали бы "
    "либо потоков, либо отдельного event loop, что усложняет код."
)
bullet(
    "Автоматическая валидация запросов через Pydantic — параметры форм, "
    "query-параметры и тела запросов проверяются декларативно, без ручных if-проверок."
)
bullet(
    "Dependency Injection (DI) из коробки — зависимости (сессия БД, текущий "
    "пользователь, роль) внедряются через Depends(), что даёт чистое разделение слоёв "
    "без глобального состояния."
)
bullet(
    "Встроенная документация OpenAPI /api/docs — при разработке и отладке не нужен "
    "Postman: все эндпоинты доступны прямо в браузере."
)
bullet(
    "Jinja2 шаблоны — для server-side rendering HTML. Отдельный JS-фронтенд "
    "(React/Vue) избыточен для внутренней корпоративной системы: увеличил бы сложность "
    "деплоя и количество кода без реального UX-выигрыша."
)

h2("2.2 ORM — SQLAlchemy 2.0")
para("SQLAlchemy 2.0 (asyncio-режим) выбран вместо Tortoise ORM, GINO или чистого asyncpg:")
bullet(
    "Зрелость и экосистема: SQLAlchemy — стандарт de-facto в Python. "
    "Документация, Stack Overflow, подключаемые инструменты (alembic, pytest-sqlalchemy) "
    "гораздо богаче альтернатив."
)
bullet(
    "Mapped-аннотации (новый синтаксис 2.0) дают строгую типизацию: IDE и mypy "
    "понимают типы колонок без runtime-сюрпризов."
)
bullet(
    "expire_on_commit=False — после commit() объекты не теряют данные. "
    "Критично в async-контексте, где N+1 ленивых запросов невозможны: "
    "нужно явно указывать selectinload() что улучшает производительность и предсказуемость."
)
bullet(
    "pool_pre_ping=True — перед выдачей соединения из пула проверяется, что "
    "PostgreSQL всё ещё жив. Без этого после простоя приложение падает с "
    "\"connection reset\" до перезапуска."
)

h2("2.3 База данных — PostgreSQL + asyncpg")
para("PostgreSQL выбран как СУБД:")
bullet(
    "ENUM-типы уровня БД: ProgramLevel (bachelor/specialist/master/postgraduate) и "
    "ReportStatusEnum (draft/submitted/approved) хранятся как SQL ENUM, "
    "а не VARCHAR. Это гарантирует целостность данных на уровне БД — "
    "приложение не может записать невалидное состояние даже через прямой SQL."
)
bullet(
    "Транзакционный DDL: Alembic-миграции выполняются в транзакции — "
    "при ошибке миграция полностью откатывается, БД не остаётся в полусломанном состоянии."
)
bullet(
    "asyncpg — самый быстрый асинхронный драйвер для PostgreSQL в Python, "
    "реализует нативный бинарный протокол без DBAPI-обёртки."
)
bullet(
    "SQLite не рассматривался: нет встроенных ENUM, слабая конкурентность записи, "
    "нет поддержки timezone-aware datetime на уровне хранения."
)

h2("2.4 Миграции — Alembic")
bullet(
    "Alembic — официальный инструмент миграций от авторов SQLAlchemy. "
    "Версионированные скрипты в alembic/versions/ хранятся в git вместе с кодом, "
    "изменения схемы воспроизводимы на любом окружении."
)
bullet(
    "Каждая миграция имеет upgrade() и downgrade() — можно откатиться назад "
    "без потери данных. Это важно при hotfix-деплоях."
)

h2("2.5 Аутентификация — JWT в cookie")
bullet(
    "JWT (HS256) хранится в HTTP cookie (не в localStorage). "
    "Cookie автоматически отправляется браузером — не нужен JS для подстановки "
    "заголовка Authorization. Работает с обычными HTML-формами (POST без fetch)."
)
bullet(
    "Срок жизни токена — 8 часов (ACCESS_TOKEN_EXPIRE_MINUTES=480). "
    "Оптимально для рабочего дня: не надо перелогиниваться, но токен "
    "не живёт неделями."
)
bullet(
    "python-jose + passlib[bcrypt]: jose — для JWT, bcrypt — для хэширования паролей. "
    "bcrypt специально спроектирован медленным (cost factor), что делает брутфорс "
    "хэшей дорогим даже при утечке БД."
)

h2("2.6 Frontend — Tailwind CSS (CDN) + Jinja2")
bullet(
    "Tailwind подключается через CDN-скрипт — нет шага сборки (webpack/vite). "
    "Для внутренней системы без публичного трафика это приемлемо: "
    "производительность не критична, зато деплой тривиален."
)
bullet(
    "Server-side rendering через Jinja2: HTML генерируется на сервере и "
    "отправляется целиком. Нет проблем с SEO, гидратацией, CORS. "
    "Для корпоративной системы с небольшой аудиторией — оптимальный выбор."
)
bullet(
    "Минимум JavaScript: только toggleAddItem() и toggleFileField() — "
    "показ/скрытие inline-форм без перезагрузки страницы. "
    "Вся бизнес-логика на сервере, JS не держит состояние."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  3. АРХИТЕКТУРА
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Архитектура системы")

h2("3.1 Общая схема слоёв")
para(
    "Система построена по трёхслойной архитектуре (Layered Architecture). "
    "Каждый слой имеет строго определённую ответственность:"
)
doc.add_paragraph()

# ASCII diagram
lines = [
    "┌─────────────────────────────────────────────────────────────────┐",
    "│                        БРАУЗЕР (HTML)                           │",
    "│              GET /programs  │  POST /reports/.../submit          │",
    "└────────────────────┬────────────────────────────────────────────┘",
    "                     │ HTTP",
    "┌────────────────────▼────────────────────────────────────────────┐",
    "│                ROUTERS  (app/routers/*.py)                       │",
    "│  • Разбор HTTP-запроса (Form, UploadFile, path-param)           │",
    "│  • Проверка прав через Depends(get_head / get_current_user)     │",
    "│  • Вызов сервисного слоя                                        │",
    "│  • Формирование ответа: TemplateResponse / RedirectResponse     │",
    "└────────────────────┬────────────────────────────────────────────┘",
    "                     │",
    "┌────────────────────▼────────────────────────────────────────────┐",
    "│               SERVICES  (app/services/*.py)                      │",
    "│  • Бизнес-логика (статусные переходы, запрет само-утверждения)  │",
    "│  • Работа с файлами (UUID, mime-проверка, сохранение на диск)   │",
    "│  • SQL-запросы к ORM                                            │",
    "└────────────────────┬────────────────────────────────────────────┘",
    "                     │",
    "┌────────────────────▼────────────────────────────────────────────┐",
    "│               MODELS  (app/models/*.py)                          │",
    "│  SQLAlchemy ORM — маппинг таблиц на Python-классы               │",
    "│  Отношения (FK, relationship, cascade)                           │",
    "└────────────────────┬────────────────────────────────────────────┘",
    "                     │  asyncpg",
    "┌────────────────────▼────────────────────────────────────────────┐",
    "│                   PostgreSQL                                     │",
    "└─────────────────────────────────────────────────────────────────┘",
]
for line in lines:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(line)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
doc.add_paragraph()

h2("3.2 Схема базы данных")
para("Таблицы и их связи:")
doc.add_paragraph()
db_lines = [
    "roles ──────────────────────────── users",
    "  id PK                              id PK",
    "  name                               full_name",
    "  display_name                       email  UNIQUE",
    "                                     hashed_password",
    "departments ──────────────────┐      role_id  FK→roles.id",
    "  id PK                       └────► department_id  FK→departments.id",
    "  name",
    "",
    "educational_programs                checklist_categories",
    "  id PK                               id PK",
    "  name / code                         name",
    "  level  ENUM                         discipline_id  FK→disciplines.id",
    "  start_year                          created_by_id  FK→users.id",
    "  director_id  FK→users.id",
    "  department_id  FK→departments.id   checklist_items",
    "                                       id PK",
    "disciplines                            title",
    "  id PK                               is_done  BOOL",
    "  name / code                          comment",
    "  goals / objectives / outcomes        category_id  FK→checklist_categories.id",
    "  program_id  FK→educational_programs",
    "  department_id  FK→departments.id   report_statuses",
    "  teacher_primary_id  FK→users.id      id PK",
    "  teacher_secondary_id  FK→users.id    discipline_id  FK  UNIQUE",
    "                                        status  ENUM  draft|submitted|approved",
    "report_files                            submitted_by_id  FK→users.id",
    "  id PK                                approved_by_id  FK→users.id",
    "  discipline_id  FK→disciplines.id      submitted_at / approved_at",
    "  category_id  FK→checklist_categories (SET NULL on delete)",
    "  original_name / stored_name",
    "  content_type / size",
    "  uploaded_by_id  FK→users.id",
    "  feedback",
]
for line in db_lines:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(line)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
doc.add_paragraph()

h2("3.3 Структура файлов проекта")
code_block([
    "edu_system/",
    "├── app/",
    "│   ├── main.py            # точка входа FastAPI, монтирование роутеров",
    "│   ├── config.py          # настройки через pydantic-settings (.env)",
    "│   ├── database.py        # движок, сессия, Base",
    "│   ├── auth.py            # JWT, хэширование, get_current_user, require_role",
    "│   ├── models/            # SQLAlchemy ORM-модели",
    "│   │   ├── __init__.py    # единый импорт — Alembic видит все таблицы",
    "│   │   ├── user.py",
    "│   │   ├── program.py",
    "│   │   ├── discipline.py",
    "│   │   ├── checklist.py",
    "│   │   ├── report.py",
    "│   │   └── report_file.py",
    "│   ├── routers/           # HTTP-слой",
    "│   │   ├── auth.py",
    "│   │   ├── programs.py",
    "│   │   ├── disciplines.py",
    "│   │   └── reports.py",
    "│   ├── services/          # бизнес-логика",
    "│   │   ├── programs.py",
    "│   │   ├── disciplines.py",
    "│   │   └── reports.py",
    "│   ├── static/",
    "│   │   └── uploads/       # загружённые файлы отчётов",
    "│   └── templates/         # Jinja2 HTML",
    "│       ├── base.html",
    "│       ├── programs/",
    "│       ├── disciplines/",
    "│       └── reports/",
    "├── alembic/               # миграции БД",
    "│   └── versions/",
    "├── seed.py                # начальные данные",
    "├── seed_teacher_reports.py",
    "├── requirements.txt",
    "└── alembic.ini",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  4. КЛЮЧЕВЫЕ РЕШЕНИЯ И РАЗБОР КОДА
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Ключевые архитектурные решения и разбор кода")

# ─── 4.1 Dependency Injection ────────────────────────────────────────────────
h2("4.1 Dependency Injection: авторизация и роли")
para(
    "Файл app/auth.py реализует всю систему авторизации через фабрику зависимостей:"
)
code_block([
    "def require_role(*role_names: str):",
    "    async def dependency(current_user: User = Depends(get_current_user)) -> User:",
    "        if current_user.role.name not in role_names:",
    "            raise HTTPException(status_code=403, detail='Недостаточно прав')",
    "        return current_user",
    "    return dependency",
    "",
    "get_head = require_role('head_of_department')",
])
para(
    "Почему именно так:"
)
bullet(
    "require_role — фабрика, возвращающая замыкание. Новая роль добавляется одной "
    "строкой: get_dean = require_role('dean'). Не нужно дублировать проверочный код."
)
bullet(
    "get_head используется как Depends(get_head) в роутерах. FastAPI сам вызывает "
    "dependency, поднимает 403 до выполнения бизнес-логики и передаёт "
    "current_user в хэндлер."
)
bullet(
    "*role_names: str — variadic, можно перечислить несколько ролей: "
    "require_role('head_of_department', 'dean') для эндпоинтов, доступных обоим."
)

# ─── 4.2 Сессия БД ───────────────────────────────────────────────────────────
h2("4.2 Управление сессией базы данных")
code_block([
    "# app/database.py",
    "engine = create_async_engine(",
    "    settings.DATABASE_URL,",
    "    echo=False,",
    "    pool_pre_ping=True,   # ← проверка живости соединения перед выдачей из пула",
    ")",
    "",
    "AsyncSessionLocal = async_sessionmaker(",
    "    bind=engine,",
    "    class_=AsyncSession,",
    "    expire_on_commit=False,  # ← объекты не «протухают» после commit()",
    ")",
    "",
    "async def get_db() -> AsyncSession:",
    "    async with AsyncSessionLocal() as session:",
    "        yield session          # ← гарантированное закрытие через context manager",
])
para("Почему expire_on_commit=False:")
bullet(
    "По умолчанию SQLAlchemy после commit() помечает все атрибуты объекта как "
    "'expired' — следующее обращение к ним делает новый SELECT. В async-контексте "
    "это невозможно без явного await refresh(). expire_on_commit=False отключает "
    "это поведение: данные остаются доступны после commit() без дополнительных запросов."
)
bullet(
    "yield session + async with — гарантирует, что сессия закроется даже при "
    "исключении в хэндлере. FastAPI вызывает генератор до конца после отправки ответа."
)

# ─── 4.3 Статусная машина отчёта ─────────────────────────────────────────────
h2("4.3 Статусная машина отчётов и запрет само-утверждения")
code_block([
    "# app/services/reports.py",
    "async def approve_report(db, discipline_id, approved_by_id) -> ReportStatus:",
    "    rs = await get_or_create_report_status(db, discipline_id)",
    "    if rs.status != ReportStatusEnum.submitted:",
    "        raise HTTPException(400, 'Отчёт должен быть отправлен на утверждение')",
    "    if rs.submitted_by_id == approved_by_id:",
    "        raise HTTPException(403, 'Нельзя утвердить собственный отчёт')",
    "    rs.status = ReportStatusEnum.approved",
    "    rs.approved_at = datetime.now(timezone.utc)",
    "    rs.approved_by_id = approved_by_id",
    "    await db.commit()",
    "    return rs",
])
para("Почему именно так:")
bullet(
    "Проверка статуса на уровне сервиса, а не только в шаблоне. "
    "Шаблон скрывает кнопку «Утвердить» для своих отчётов, но это лишь UX. "
    "Без серверной проверки прямой POST на /approve всё равно бы прошёл."
)
bullet(
    "submitted_by_id хранится в момент подачи, а не вычисляется по логу событий. "
    "Это делает проверку O(1) — один SELECT, без JOIN истории."
)
bullet(
    "Двойная защита (frontend + backend) — стандартная практика defense-in-depth: "
    "UI даёт понятный отклик, сервер — гарантию."
)
bullet(
    "Переходы разрешены только вперёд (draft→submitted→approved), "
    "откат — только через явный /reset. Невалидные переходы отбиваются 400."
)

# ─── 4.4 Загрузка файлов ─────────────────────────────────────────────────────
h2("4.4 Загрузка и хранение файлов отчётов")
code_block([
    "# app/services/reports.py",
    "ALLOWED_TYPES = {",
    "    'application/pdf',",
    "    'application/msword',",
    "    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',",
    "}",
    "MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 МБ",
    "",
    "async def add_report_file(db, discipline_id, uploaded_by_id, file, category_id=None):",
    "    content = await file.read()",
    "    if len(content) > MAX_FILE_SIZE:",
    "        raise HTTPException(400, 'Файл слишком большой')",
    "    if file.content_type not in ALLOWED_TYPES:",
    "        raise HTTPException(400, 'Разрешены только PDF, DOC, DOCX')",
    "    stored_name = f'{uuid.uuid4().hex}{suffix}'",
    "    (UPLOAD_DIR / stored_name).write_bytes(content)",
    "    rf = ReportFile(..., category_id=category_id)",
])
para("Почему именно так:")
bullet(
    "UUID в имени файла на диске (stored_name). Оригинальное имя хранится в БД "
    "(original_name). Это решает: конфликты имён (два файла 'report.pdf' "
    "не перезапишут друг друга), невозможность path-traversal атаки "
    "через имя файла, удобное скачивание с правильным именем через FileResponse."
)
bullet(
    "content_type проверяется по MIME-типу из заголовка, а не по расширению. "
    "Расширение легко подделать — MIME сложнее."
)
bullet(
    "category_id=None — nullable FK с ON DELETE SET NULL. "
    "Если категория удаляется, файл не удаляется вместе с ней — "
    "данные не теряются, файл просто становится 'без категории'."
)

# ─── 4.5 Модели и relationship ───────────────────────────────────────────────
h2("4.5 ORM-модели: Mapped-аннотации и каскады")
code_block([
    "# app/models/discipline.py",
    "class Discipline(Base):",
    "    checklist_categories: Mapped[list['ChecklistCategory']] = relationship(",
    "        'ChecklistCategory', back_populates='discipline',",
    "        cascade='all, delete-orphan'   # ← каскадное удаление",
    "    )",
    "    report_status: Mapped[Optional['ReportStatus']] = relationship(",
    "        'ReportStatus', uselist=False, cascade='all, delete-orphan'",
    "    )",
])
para("Почему cascade='all, delete-orphan':")
bullet(
    "При удалении дисциплины автоматически удаляются все её категории, пункты "
    "и статус отчёта. Без каскада образовались бы 'осиротевшие' строки, "
    "нарушающие FK-ограничения."
)
bullet(
    "uselist=False для report_status — отношение 1:1 (один статус на дисциплину), "
    "SQLAlchemy возвращает объект, а не список."
)
code_block([
    "# app/models/checklist.py",
    "class ChecklistCategory(Base):",
    "    items: Mapped[list['ChecklistItem']] = relationship(",
    "        'ChecklistItem', back_populates='category',",
    "        cascade='all, delete-orphan',",
    "        order_by='ChecklistItem.id'   # ← стабильный порядок без ORDER BY в запросе",
    "    )",
    "    files: Mapped[list['ReportFile']] = relationship(",
    "        'ReportFile', back_populates='category',",
    "        order_by='ReportFile.id'",
    "    )",
])
bullet(
    "order_by='ChecklistItem.id' внутри relationship — элементы чеклиста "
    "всегда идут в порядке добавления. Альтернатива — ORDER BY в каждом запросе, "
    "что дублирует логику."
)

# ─── 4.6 selectinload ────────────────────────────────────────────────────────
h2("4.6 Загрузка связей: selectinload vs lazy loading")
code_block([
    "# app/services/reports.py",
    "result = await db.execute(",
    "    select(ChecklistCategory)",
    "    .options(",
    "        selectinload(ChecklistCategory.items),",
    "        selectinload(ChecklistCategory.files),  # ← явная загрузка файлов",
    "    )",
    "    .where(ChecklistCategory.discipline_id == discipline_id)",
    ")",
])
para("Почему selectinload, а не joinedload или lazy:")
bullet(
    "Lazy loading (по умолчанию) в async-SQLAlchemy вызвал бы MissingGreenlet-ошибку "
    "при обращении к связи вне event loop. В async-коде lazy load запрещён."
)
bullet(
    "joinedload делает JOIN в том же запросе — эффективен для отношений to-one. "
    "Для to-many (items, files) он дублирует строки родителя в результате, "
    "что увеличивает трафик."
)
bullet(
    "selectinload делает отдельный SELECT ... WHERE id IN (...) — "
    "один запрос на все дочерние объекты, без дублирования. "
    "Оптимален для коллекций."
)

# ─── 4.7 Конфиг ──────────────────────────────────────────────────────────────
h2("4.7 Конфигурация через pydantic-settings")
code_block([
    "# app/config.py",
    "class Settings(BaseSettings):",
    "    model_config = SettingsConfigDict(env_file='.env', env_file_encoding='utf-8')",
    "",
    "    DATABASE_URL: str = 'postgresql+asyncpg://...'",
    "    SECRET_KEY: str = 'change-me-in-production...'",
    "    ACCESS_TOKEN_EXPIRE_MINUTES: int = 480",
    "    ALGORITHM: str = 'HS256'",
    "",
    "settings = Settings()",
])
para("Почему pydantic-settings:")
bullet(
    "Приоритет: переменные окружения переопределяют .env-файл. "
    "В production достаточно выставить DATABASE_URL и SECRET_KEY в env — "
    "файл .env для локальной разработки."
)
bullet(
    "Типизация и валидация: если DATABASE_URL не задан — явная ошибка "
    "при старте, а не cryptic ошибка при первом запросе к БД."
)
bullet(
    "Singleton settings — единственный экземпляр создаётся при импорте, "
    "все модули используют один и тот же объект."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  5. МАРШРУТ УТВЕРЖДЕНИЯ ОТЧЁТА
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Бизнес-процесс: маршрут утверждения отчёта")

para("Жизненный цикл отчёта по дисциплине:")
doc.add_paragraph()
flow_lines = [
    "  Преподаватель                    Заведующий кафедрой",
    "       │                                    │",
    "       │  1. Создаёт раздел                 │",
    "       │     (ChecklistCategory)            │",
    "       │  2. Прикрепляет файл               │",
    "       │     (ReportFile)                   │",
    "       │  3. Заполняет чеклист              │",
    "       │     (ChecklistItem.is_done=True)   │",
    "       │  4. Нажимает «Отправить»           │",
    "       │     status: draft→submitted        │",
    "       │     submitted_by_id = teacher.id   │",
    "       │────────────────────────────────────►│",
    "       │                         5. Видит «Утвердить»",
    "       │                            (т.к. submitted_by_id ≠ head.id)",
    "       │                         6. Нажимает «Утвердить»",
    "       │                            status: submitted→approved",
    "       │                            approved_by_id = head.id",
    "       │◄────────────────────────────────────│",
    "       │                         ИЛИ",
    "       │                         6b. Пишет замечание к файлу",
    "       │                             (ReportFile.feedback)",
    "       │                         6c. Нажимает «Вернуть»",
    "       │                             status: any→draft",
    "       │◄────────────────────────────────────│",
    "       │  7. Исправляет, повторяет           │",
]
for line in flow_lines:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(line)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
doc.add_paragraph()

para("Ограничение само-утверждения:")
para(
    "Заведующий кафедры тоже является пользователем системы и теоретически может "
    "создать отчёт по дисциплине, которую ведёт лично. В этом случае система "
    "показывает информационный блок «Ожидает утверждения вышестоящим руководством» "
    "вместо кнопки «Утвердить». На backend-уровне сервис вернёт HTTP 403 при попытке "
    "самоутверждения через прямой POST-запрос."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  6. ПРИНЯТЫЕ РЕШЕНИЯ И АЛЬТЕРНАТИВЫ
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Принятые решения и рассмотренные альтернативы")

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"

hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Аспект"
hdr_cells[1].text = "Принятое решение"
hdr_cells[2].text = "Почему не альтернатива"

for cell in hdr_cells:
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ("Web-фреймворк",       "FastAPI",             "Django — синхронный ORM, избыточен;\nFlask — нет DI, нет async из коробки"),
    ("ORM",                 "SQLAlchemy 2.0 async","Tortoise ORM — меньше экосистемы;\nGINO — заброшен"),
    ("СУБД",                "PostgreSQL",           "SQLite — нет ENUM, слабая конкурентность;\nMySQL — хуже transactional DDL"),
    ("Аутентификация",      "JWT в cookie",        "Session в Redis — доп. инфра;\nlocalStorage JWT — нужен JS для каждого запроса"),
    ("Frontend",            "Jinja2 SSR + Tailwind","React/Vue — избыточно для корпоратива;\nBootstrap — больший CSS-бандл без утилитарного подхода"),
    ("Хранение файлов",     "Локальная FS (UUID)",  "S3/MinIO — доп. инфра, избыточно для одного узла;\nDBO хранение — не масштабируется"),
    ("Миграции",            "Alembic",              "django-migrate — привязан к Django;\nFlyway — Java-экосистема"),
    ("Хэширование паролей", "bcrypt (passlib)",     "SHA-256 — не предназначен для паролей;\nargon2 — сложнее настройка cost-factor"),
]

for aspect, decision, alt in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = aspect
    row_cells[1].text = decision
    row_cells[2].text = alt
    for cell in row_cells:
        for para_obj in cell.paragraphs:
            for run_obj in para_obj.runs:
                run_obj.font.name = "Arial"
                run_obj.font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  7. БЕЗОПАСНОСТЬ
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Безопасность")

h2("7.1 Что реализовано")
bullet("Пароли хранятся как bcrypt-хэш (cost factor ~12), никогда в открытом виде.")
bullet("JWT подписан HS256, секрет хранится в .env, не в коде.")
bullet("Авторизация проверяется на уровне сервиса (двойная защита: UI + backend).")
bullet("Загружаемые файлы: UUID-имя на диске, проверка MIME, ограничение 20 МБ.")
bullet("Redirect после POST (PRG-паттерн) — предотвращает двойную отправку формы при F5.")

h2("7.2 Что стоит добавить в production")
bullet("HTTPS (TLS termination на nginx/caddy перед uvicorn).")
bullet("HttpOnly + Secure флаги на cookie с JWT.")
bullet("Rate limiting на /auth/login (например, slowapi).")
bullet("Валидация MIME не только по заголовку, но и по magic bytes файла (python-magic).")
bullet("Хранение файлов вне директории static (чтобы веб-сервер не раздавал их напрямую).")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  8. ЗАПУСК
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Запуск и конфигурация")

code_block([
    "# 1. Создать виртуальное окружение и установить зависимости",
    "python3 -m venv venv && source venv/bin/activate",
    "pip install -r requirements.txt",
    "",
    "# 2. Создать .env (или выставить переменные окружения)",
    "DATABASE_URL=postgresql+asyncpg://user:pass@localhost:5432/edu_system",
    "SECRET_KEY=your-32-char-secret",
    "",
    "# 3. Применить миграции",
    "alembic upgrade head",
    "",
    "# 4. Загрузить начальные данные",
    "python3 seed.py",
    "python3 seed_teacher_reports.py   # тестовые отчёты от преподавателей",
    "",
    "# 5. Запустить сервер",
    "uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload",
    "",
    "# Аккаунты после seed:",
    "# head@vgtu.ru / password123   — заведующий кафедрой",
    "# teacher1@vgtu.ru / password123",
    "# teacher2@vgtu.ru / password123",
])

doc.add_paragraph()
para(
    "Документация API (OpenAPI / Swagger) доступна по адресу http://localhost:8000/api/docs "
    "после запуска сервера.",
    italic=True
)

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
