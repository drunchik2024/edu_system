"""Скрипт генерации документа «Описание программы» в формате docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def set_heading_color(paragraph, color: RGBColor):
    for run in paragraph.runs:
        run.font.color.rgb = color


def add_code_block(doc, code: str, lang_comment: str = ""):
    """Добавляет блок кода с моноширинным шрифтом и серым фоном."""
    if lang_comment:
        p = doc.add_paragraph()
        run = p.add_run(f"# {lang_comment}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.name = "Courier New"

    for line in code.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        # серый фон через shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F3F4F6")
        pPr.append(shd)


def add_note(doc, text: str):
    """Добавляет примечание курсивом."""
    p = doc.add_paragraph()
    run = p.add_run(f"Примечание: {text}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def main():
    doc = Document()

    # ──────── Стили ────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    for h_name in ["Heading 1", "Heading 2", "Heading 3"]:
        s = doc.styles[h_name]
        s.font.name = "Arial"
        s.font.bold = True

    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 3"].font.size = Pt(12)

    # ──────── Титул ────────
    title = doc.add_heading("ИС Управления Образовательными Программами — ВГТУ", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Полное техническое описание системы")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    # ══════════════════════════════════════════
    # 1. О ЧЁМ ЭТА СИСТЕМА
    # ══════════════════════════════════════════
    doc.add_heading("1. О чём эта система и зачем она нужна", 1)

    doc.add_paragraph(
        "Система предназначена для заведующих кафедрами Воронежского государственного "
        "технического университета (ВГТУ). Её задача — дать заведующему единое рабочее место, "
        "где он видит все образовательные программы своей кафедры, дисциплины внутри них, "
        "может отслеживать статус отчётов по каждой дисциплине и утверждать их."
    )

    doc.add_paragraph("Что конкретно умеет система:")
    for item in [
        "Показывать список образовательных программ кафедры с фильтрацией по названию, "
        "коду и ФИО руководителя программы.",
        "По каждой программе — две вкладки: «Дисциплины кафедры» (ведёт наша кафедра) "
        "и «Другие кафедры» (дисциплины программы, которые ведут чужие кафедры).",
        "По каждой дисциплине — страница с чеклистом: категории → пункты → галочки. "
        "К категории можно прикрепить PDF/Word-файлы.",
        "Управлять жизненным циклом отчёта: черновик → на утверждении → утверждено. "
        "Заведующий, подавая отчёт сам за себя, сразу получает статус «Утверждено» "
        "(минует промежуточный шаг). Откат в черновик тоже доступен.",
        "Запрет само-утверждения: один заведующий не может утвердить отчёт, который сам же подал.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    # ══════════════════════════════════════════
    # 2. СТЕК ТЕХНОЛОГИЙ
    # ══════════════════════════════════════════
    doc.add_heading("2. Стек технологий и почему именно он", 1)

    # 2.1 Python
    doc.add_heading("2.1 Python 3.12", 2)
    doc.add_paragraph(
        "Язык реализации — Python 3.12. Главная причина — нативная поддержка синтаксиса "
        "«X | None» вместо «Optional[X]» и улучшенная производительность asyncio. "
        "Python широко используется в академической среде, поэтому проект проще поддерживать "
        "людям без глубокого коммерческого опыта."
    )

    # 2.2 FastAPI
    doc.add_heading("2.2 FastAPI", 2)
    doc.add_paragraph(
        "FastAPI — современный веб-фреймворк на Python. Выбрали его вместо Flask/Django по трём причинам:"
    )
    for item in [
        "Async из коробки. FastAPI строится на Starlette + asyncio: все обработчики запросов "
        "могут быть async def, что позволяет ждать результата из базы без блокировки потока.",
        "Dependency Injection (DI). Зависимости (БД-сессия, текущий пользователь) "
        "описываются декларативно через Depends(). Это чище, чем глобальные объекты Flask.",
        "Автодокументация. Достаточно аннотаций типов — FastAPI сам генерирует Swagger UI "
        "на /api/docs. При разработке это экономит время.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_paragraph("Пример маршрута с DI:")
    add_code_block(doc, """\
@router.get("/{program_id}", response_class=HTMLResponse)
async def program_detail(
    program_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),          # сессия БД
    current_user: User = Depends(get_current_user),  # проверка JWT
):
    program = await svc_programs.get_program_by_id(db, program_id, ...)
    ...""", "routers/programs.py")

    doc.add_paragraph(
        "FastAPI сам вызовет get_db() и get_current_user() перед тем, как выполнить тело "
        "функции. Если токен невалиден — вернёт 401 ещё до входа в обработчик."
    )

    # 2.3 SQLAlchemy
    doc.add_heading("2.3 SQLAlchemy 2.0 (async ORM)", 2)
    doc.add_paragraph(
        "ORM (Object-Relational Mapper) — прослойка между Python-объектами и таблицами БД. "
        "Вместо написания SQL-запросов вручную, описываем модели как классы:"
    )
    add_code_block(doc, """\
class Discipline(Base):
    __tablename__ = "disciplines"

    id: Mapped[int] = mapped_column(primary_key=True)
    name: Mapped[str] = mapped_column(String(512))
    program_id: Mapped[int] = mapped_column(ForeignKey("educational_programs.id"))

    program: Mapped["EducationalProgram"] = relationship(
        "EducationalProgram", back_populates="disciplines"
    )""", "models/discipline.py")

    doc.add_paragraph(
        "SQLAlchemy 2.0 выбран (а не 1.4) из-за нового типизированного API «Mapped[T]» — "
        "IDE видит типы полей и подсказывает ошибки ещё до запуска. "
        "Async-версия (AsyncSession, create_async_engine) нужна, чтобы запросы к PostgreSQL "
        "не блокировали event loop FastAPI."
    )
    add_note(doc, "Почему не «чистый» asyncpg без ORM? ORM даёт типобезопасность, "
             "кэш отношений и selectinload. Без ORM пришлось бы писать JOIN-ы и маппинг вручную.")

    # 2.4 asyncpg
    doc.add_heading("2.4 asyncpg", 2)
    doc.add_paragraph(
        "asyncpg — самый быстрый async-драйвер для PostgreSQL на Python (написан на Cython). "
        "SQLAlchemy использует его «под капотом» через строку подключения: "
        "postgresql+asyncpg://... Нам не нужно вызывать asyncpg напрямую — "
        "это просто транспортный слой."
    )

    # 2.5 PostgreSQL
    doc.add_heading("2.5 PostgreSQL", 2)
    doc.add_paragraph(
        "В качестве СУБД выбрали PostgreSQL (а не SQLite или MySQL) по следующим причинам:"
    )
    for item in [
        "Полная поддержка ACID-транзакций и внешних ключей с ON DELETE действиями.",
        "Нативные типы ENUM — статус отчёта хранится как типизированный перечислитель, "
        "а не просто строка. PostgreSQL не даст записать невалидный статус.",
        "Хорошая поддержка timezone-aware дат (TIMESTAMPTZ) — важно при логировании событий.",
        "SQLite не подходит для параллельных async-запросов: он блокирует файл целиком при записи.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    # 2.6 Alembic
    doc.add_heading("2.6 Alembic (миграции)", 2)
    doc.add_paragraph(
        "Alembic — инструмент миграций для SQLAlchemy. Каждое изменение схемы БД фиксируется "
        "как отдельный файл-миграция с функциями upgrade() и downgrade(). "
        "Это позволяет применять изменения последовательно на любом окружении (dev, stage, prod) "
        "и откатываться назад при проблемах."
    )
    doc.add_paragraph("Пример миграции — добавление столбца category_id к файлам отчёта:")
    add_code_block(doc, """\
def upgrade() -> None:
    op.add_column(
        "report_files",
        sa.Column(
            "category_id",
            sa.Integer(),
            sa.ForeignKey("checklist_categories.id", ondelete="SET NULL"),
            nullable=True,
        ),
    )

def downgrade() -> None:
    op.drop_column("report_files", "category_id")""", "alembic/versions/a1b2c3d4e5f6_...")

    doc.add_paragraph(
        "Ключевое решение: ondelete='SET NULL' — если категория удаляется, "
        "файл не удаляется вместе с ней, а просто «теряет» привязку к категории."
    )

    # 2.7 Jinja2 + Tailwind
    doc.add_heading("2.7 Jinja2 + Tailwind CSS", 2)
    doc.add_paragraph(
        "Рендеринг HTML — серверный (SSR, Server-Side Rendering): "
        "FastAPI вызывает Jinja2, который подставляет данные в шаблон и возвращает "
        "готовый HTML браузеру. Это проще и быстрее, чем React/Vue, когда нет нужды "
        "в богатом интерактивном интерфейсе."
    )
    doc.add_paragraph(
        "Tailwind CSS подключён через CDN (без сборщика). "
        "Стили — утилитарные классы прямо в HTML: class=\"text-sm text-gray-600 px-4 py-2\". "
        "Никакого отдельного CSS-файла писать не нужно."
    )
    add_note(doc, "Почему не React/Vue? Для внутренней CRUD-системы SSR достаточен. "
             "SPA-фреймворки добавили бы сборку, API-слой, state management — "
             "всё это лишняя сложность без видимой пользы.")

    # 2.8 JWT
    doc.add_heading("2.8 JWT-авторизация (python-jose + passlib)", 2)
    doc.add_paragraph(
        "Авторизация построена на JWT-токенах (JSON Web Token). "
        "После успешного входа сервер создаёт токен и кладёт его в HTTP-cookie:"
    )
    add_code_block(doc, """\
token = create_access_token({"sub": str(user.id)})
response.set_cookie(
    key="access_token",
    value=token,
    httponly=True,   # JS не может прочитать — защита от XSS
    max_age=60 * 60 * 8,  # 8 часов
    samesite="lax",  # защита от CSRF
)""", "routers/auth.py")

    doc.add_paragraph("При каждом запросе get_current_user() читает cookie, "
                      "декодирует JWT и возвращает объект User:")
    add_code_block(doc, """\
async def get_current_user(request: Request, db: AsyncSession = Depends(get_db)) -> User:
    token = request.cookies.get("access_token")
    payload = jwt.decode(token, settings.SECRET_KEY, algorithms=["HS256"])
    user_id = int(payload.get("sub"))
    user = await db.get(User, user_id)
    return user""", "auth.py (упрощённо)")

    doc.add_paragraph(
        "Пароли хэшируются через bcrypt (passlib). bcrypt намеренно медленный алгоритм — "
        "это защита от брутфорса при утечке базы данных: перебор 1 млн паролей занимает "
        "часы, а не секунды."
    )

    # ══════════════════════════════════════════
    # 3. АРХИТЕКТУРА
    # ══════════════════════════════════════════
    doc.add_heading("3. Архитектура системы", 1)

    doc.add_heading("3.1 Трёхслойная архитектура", 2)
    doc.add_paragraph(
        "Код разбит на три слоя с чёткими обязанностями:"
    )

    # Таблица слоёв
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Слой"
    hdr[1].text = "Файлы"
    hdr[2].text = "Что делает"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    rows_data = [
        ("Роутеры (HTTP)", "app/routers/*.py",
         "Принимают HTTP-запрос, проверяют права, вызывают сервис, возвращают HTML или редирект"),
        ("Сервисы (бизнес-логика)", "app/services/*.py",
         "Содержат бизнес-правила: логика перехода статусов, фильтрация, подсчёт. "
         "Не знают о HTTP — только о данных"),
        ("Модели (данные)", "app/models/*.py",
         "ORM-классы = таблицы БД. Описывают структуру данных и связи между таблицами"),
    ]
    for i, (layer, files, desc) in enumerate(rows_data, 1):
        row = table.rows[i].cells
        row[0].text = layer
        row[1].text = files
        row[2].text = desc

    doc.add_paragraph()
    doc.add_paragraph(
        "Почему три слоя, а не «всё в роутере»? Если в роутере написать и SQL-запрос, "
        "и бизнес-проверку, и рендеринг — при изменении логики придётся лезть в один "
        "большой файл и искать нужное место. Разделение упрощает тестирование: "
        "сервис можно вызвать напрямую без HTTP-запроса."
    )

    doc.add_heading("3.2 Структура файлов", 2)
    add_code_block(doc, """\
edu_system/
├── app/
│   ├── main.py          # точка входа FastAPI
│   ├── auth.py          # JWT, bcrypt, get_current_user
│   ├── config.py        # настройки из .env
│   ├── database.py      # движок и сессия SQLAlchemy
│   ├── models/          # ORM-модели (= таблицы БД)
│   │   ├── user.py
│   │   ├── program.py
│   │   ├── discipline.py
│   │   ├── checklist.py
│   │   ├── report.py
│   │   └── report_file.py
│   ├── routers/         # HTTP-обработчики
│   │   ├── auth.py
│   │   ├── programs.py
│   │   ├── disciplines.py
│   │   └── reports.py
│   ├── services/        # бизнес-логика
│   │   ├── programs.py
│   │   ├── disciplines.py
│   │   └── reports.py
│   ├── templates/       # Jinja2 HTML-шаблоны
│   │   ├── base.html
│   │   ├── auth/login.html
│   │   ├── programs/list.html
│   │   ├── programs/detail.html
│   │   └── reports/checklist.html
│   └── static/
│       └── uploads/     # загружённые файлы отчётов
├── alembic/             # миграции БД
│   └── versions/
├── seed.py              # тестовые данные
├── .env                 # секреты (не в git)
└── requirements.txt""", "структура проекта")

    doc.add_heading("3.3 Поток запроса (от браузера до БД)", 2)
    doc.add_paragraph(
        "Рассмотрим полный путь запроса «открыть страницу программы»:"
    )
    steps = [
        ("Браузер", "GET /programs/42 с cookie access_token=eyJ..."),
        ("FastAPI", "Находит маршрут в programs.router. Запускает DI-граф."),
        ("get_db()", "Открывает AsyncSession к PostgreSQL через пул соединений."),
        ("get_current_user()", "Читает cookie, декодирует JWT, загружает User из БД."),
        ("program_detail()", "Вызывает svc_programs.get_program_by_id() и svc_disciplines.get_disciplines_for_department()."),
        ("Сервисы", "Выполняют SELECT с JOIN/WHERE/ORDER BY, возвращают list[dict]."),
        ("Jinja2", "Рендерит programs/detail.html, подставляет данные."),
        ("Браузер", "Получает готовый HTML, отображает страницу."),
    ]
    table2 = doc.add_table(rows=len(steps) + 1, cols=2)
    table2.style = "Table Grid"
    table2.rows[0].cells[0].text = "Участник"
    table2.rows[0].cells[1].text = "Действие"
    for cell in table2.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, (actor, action) in enumerate(steps, 1):
        table2.rows[i].cells[0].text = actor
        table2.rows[i].cells[1].text = action

    doc.add_paragraph()

    # ══════════════════════════════════════════
    # 4. БАЗА ДАННЫХ
    # ══════════════════════════════════════════
    doc.add_heading("4. База данных", 1)

    doc.add_heading("4.1 Схема таблиц", 2)
    doc.add_paragraph(
        "Всего 9 таблиц. Ниже описание каждой и зачем она нужна:"
    )

    tables_info = [
        ("roles", "id, name",
         "Роли пользователей. Сейчас одна роль: head_of_department (заведующий кафедрой). "
         "Отдельная таблица (а не строка в users) — чтобы роли можно было добавлять без "
         "изменения схемы."),
        ("departments", "id, name",
         "Кафедры университета. Пользователь принадлежит кафедре — это его «рабочая зона»: "
         "он видит только программы и дисциплины своей кафедры."),
        ("users", "id, full_name, email, hashed_password, role_id, department_id",
         "Учётные записи. Пароль хранится как bcrypt-хэш, никогда в открытом виде. "
         "department_id nullable — на случай суперпользователя без привязки к кафедре."),
        ("educational_programs", "id, name, code, level, start_year, description, director_id, department_id",
         "Образовательные программы. level — ENUM: bachelor/specialist/master/postgraduate. "
         "director_id — FK на users (руководитель программы)."),
        ("disciplines", "id, name, code, goals, objectives, outcomes, program_id, department_id, teacher_primary_id, teacher_secondary_id",
         "Дисциплины внутри программ. Две FK на users: ответственный и ассистент. "
         "department_id показывает, какая кафедра ведёт дисциплину — может отличаться от "
         "кафедры программы («внешние» дисциплины)."),
        ("report_statuses", "id, discipline_id (unique), status, submitted_at, approved_at, submitted_by_id, approved_by_id",
         "Статус отчёта по дисциплине. discipline_id UNIQUE — один статус на одну дисциплину. "
         "submitted_by_id и approved_by_id — кто подал и кто утвердил."),
        ("checklist_categories", "id, name, discipline_id, created_by_id",
         "Категории чеклиста (папки). Например: «Учебно-методический комплекс», «Рабочая программа»."),
        ("checklist_items", "id, title, is_done, comment, category_id, created_by_id",
         "Пункты внутри категории. is_done — булев флаг выполнения."),
        ("report_files", "id, discipline_id, original_name, stored_name, content_type, size, uploaded_at, uploaded_by_id, feedback, category_id",
         "Загруженные файлы (PDF/DOC/DOCX). stored_name — UUID-имя на диске (исключает коллизии). "
         "category_id nullable FK → при удалении категории файл остаётся, category_id становится NULL."),
    ]

    for tname, cols, desc in tables_info:
        doc.add_heading(tname, 3)
        p = doc.add_paragraph()
        p.add_run("Столбцы: ").bold = True
        p.add_run(cols).font.name = "Courier New"
        doc.add_paragraph(desc)

    doc.add_heading("4.2 Ключевые связи (Foreign Keys)", 2)
    add_code_block(doc, """\
users ─────────────────────────────────────────────────────────────────────────────
  ↑ role_id → roles.id
  ↑ department_id → departments.id

educational_programs
  ↑ director_id → users.id
  ↑ department_id → departments.id

disciplines
  ↑ program_id → educational_programs.id
  ↑ department_id → departments.id
  ↑ teacher_primary_id → users.id
  ↑ teacher_secondary_id → users.id

report_statuses
  ↑ discipline_id (UNIQUE) → disciplines.id   ← один статус на дисциплину
  ↑ submitted_by_id → users.id
  ↑ approved_by_id → users.id

checklist_categories
  ↑ discipline_id → disciplines.id

checklist_items
  ↑ category_id → checklist_categories.id  (CASCADE DELETE)

report_files
  ↑ discipline_id → disciplines.id  (CASCADE DELETE)
  ↑ category_id → checklist_categories.id  (ON DELETE SET NULL)""", "схема связей")

    doc.add_paragraph(
        "cascade='all, delete-orphan' в ORM — это «мягкий» каскад на уровне Python, "
        "который работает вместе с ON DELETE CASCADE в PostgreSQL. "
        "Удалил дисциплину — автоматически удалятся её категории, пункты и файлы."
    )

    # ══════════════════════════════════════════
    # 5. АВТОРИЗАЦИЯ И РОЛИ
    # ══════════════════════════════════════════
    doc.add_heading("5. Авторизация и ролевая модель", 1)

    doc.add_heading("5.1 Как работает вход", 2)
    doc.add_paragraph("Процесс входа в систему:")
    for step in [
        "Пользователь отправляет форму (email + пароль) на POST /auth/login.",
        "Сервер ищет пользователя в БД по email.",
        "Проверяет пароль через bcrypt: verify_password(plain, hashed).",
        "При успехе создаёт JWT-токен: payload = {\"sub\": str(user.id), \"exp\": ...}.",
        "Кладёт токен в httponly-cookie. Делает редирект на /programs.",
        "Все последующие запросы автоматически несут cookie — браузер добавляет её сам.",
    ]:
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)

    doc.add_heading("5.2 Фабрика ролей require_role()", 2)
    doc.add_paragraph(
        "Вместо проверки роли вручную в каждом обработчике, используется фабрика зависимостей:"
    )
    add_code_block(doc, """\
def require_role(*role_names: str):
    async def dependency(current_user: User = Depends(get_current_user)) -> User:
        if current_user.role.name not in role_names:
            raise HTTPException(status_code=403, detail="Недостаточно прав")
        return current_user
    return dependency

# Готовая зависимость для заведующего кафедрой:
get_head = require_role("head_of_department")

# Использование в роутере:
@router.post("/discipline/{discipline_id}/submit")
async def submit_report(
    discipline_id: int,
    current_user: User = Depends(get_current_user),  # любой авторизованный
):
    ...

@router.post("/discipline/{discipline_id}/approve")
async def approve_report(
    discipline_id: int,
    current_user: User = Depends(get_head),  # только заведующий
):
    ...""", "auth.py + routers/reports.py")

    doc.add_heading("5.3 Запрет само-утверждения", 2)
    doc.add_paragraph(
        "Бизнес-правило: один и тот же заведующий не может подать отчёт и сам же его утвердить. "
        "Это защита от злоупотреблений. Реализовано на двух уровнях:"
    )
    add_code_block(doc, """\
# Уровень 1 — бэкенд (services/reports.py):
async def approve_report(db, discipline_id, approved_by_id):
    rs = await get_or_create_report_status(db, discipline_id)
    if rs.submitted_by_id == approved_by_id:
        raise HTTPException(status_code=403, detail="Нельзя утвердить собственный отчёт")
    ...

# Уровень 2 — фронтенд (reports/checklist.html):
{% if report_status.submitted_by_id == user.id %}
  <p class="text-yellow-700">Вы подали этот отчёт — утверждение недоступно</p>
{% else %}
  <form method="post" action=".../approve">
    <button>Утвердить</button>
  </form>
{% endif %}""", "двухуровневая защита")

    add_note(doc, "Фронтенд скрывает кнопку для удобства UX, "
             "но бэкенд тоже проверяет — на случай прямого HTTP-запроса.")

    # ══════════════════════════════════════════
    # 6. ЖИЗНЕННЫЙ ЦИКЛ ОТЧЁТА
    # ══════════════════════════════════════════
    doc.add_heading("6. Жизненный цикл отчёта (машина состояний)", 1)

    doc.add_paragraph("Каждая дисциплина имеет один отчёт со статусом:")
    add_code_block(doc, """\
                   ┌────────────────────────────────────────────────┐
                   │                                                │
               [черновик]                                      [сброс]
                   │                                                │
             подать отчёт                                    (reset)
                   │                                                │
                   ▼                                                │
         [на утверждении]  ──── утвердить (другой зав.) ────▶  [утверждено]
                   │
                   └──── (если подаёт зав. кафедры) ──────▶  [утверждено]
                         auto_approve = True""", "схема переходов")

    doc.add_paragraph("Код машины состояний:")
    add_code_block(doc, """\
async def submit_report(db, discipline_id, submitted_by_id, auto_approve=False):
    rs = await get_or_create_report_status(db, discipline_id)

    if rs.status == ReportStatusEnum.approved:
        raise HTTPException(400, "Отчёт уже утверждён")

    rs.submitted_at = datetime.now(timezone.utc)
    rs.submitted_by_id = submitted_by_id

    if auto_approve:
        # Заведующий кафедры: сразу утверждено
        rs.status = ReportStatusEnum.approved
        rs.approved_at = rs.submitted_at
        rs.approved_by_id = submitted_by_id
    else:
        rs.status = ReportStatusEnum.submitted  # ждёт другого зав.

    await db.commit()""", "services/reports.py")

    doc.add_paragraph("Флаг auto_approve устанавливается в роутере на основе роли:")
    add_code_block(doc, """\
@router.post("/discipline/{discipline_id}/submit")
async def submit_report(discipline_id, current_user = Depends(get_current_user), ...):
    auto_approve = current_user.role.name == "head_of_department"
    await svc.submit_report(db, discipline_id, current_user.id, auto_approve=auto_approve)""",
                   "routers/reports.py")

    add_note(doc, "Роль-логика (auto_approve) живёт в роутере, а не в сервисе. "
             "Сервис «не знает» о ролях — он просто выполняет переданную инструкцию. "
             "Это разделение ответственности.")

    # ══════════════════════════════════════════
    # 7. ЗАГРУЗКА ФАЙЛОВ
    # ══════════════════════════════════════════
    doc.add_heading("7. Загрузка и хранение файлов", 1)

    doc.add_heading("7.1 Как работает загрузка", 2)
    for step in [
        "Форма с enctype='multipart/form-data' отправляет файл на POST /reports/discipline/{id}/category.",
        "FastAPI принимает UploadFile (streaming upload, не грузит всё в память сразу).",
        "Читаем содержимое, проверяем размер (≤ 20 МБ) и MIME-тип (только PDF/DOC/DOCX).",
        "Генерируем UUID-имя: f\"{uuid4().hex}{suffix}\" — это исключает коллизии и path-traversal атаки.",
        "Записываем файл в app/static/uploads/.",
        "Сохраняем запись в таблицу report_files (оригинальное имя, UUID-имя, размер, тип).",
    ]:
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)

    doc.add_heading("7.2 Безопасность", 2)
    add_code_block(doc, """\
ALLOWED_TYPES = {
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 МБ

content = await file.read()
if len(content) > MAX_FILE_SIZE:
    raise HTTPException(400, "Файл слишком большой (макс. 20 МБ)")
if file.content_type not in ALLOWED_TYPES:
    raise HTTPException(400, "Разрешены только PDF, DOC, DOCX")

# UUID-имя — оригинальное имя файла не попадает в путь на диске
stored_name = f"{uuid4().hex}{suffix}"
(UPLOAD_DIR / stored_name).write_bytes(content)""", "services/reports.py")

    doc.add_paragraph("Почему UUID-имя? Если пользователь загрузит файл "
                      "../../etc/passwd — UUID исключит path-traversal: "
                      "имя не участвует в путях, только генерированный идентификатор.")

    doc.add_heading("7.3 Скачивание файла", 2)
    add_code_block(doc, """\
@router.get("/files/{file_id}/download")
async def download_file(file_id: int, ...):
    rf = await db.get(ReportFile, file_id)
    path = UPLOAD_DIR / rf.stored_name  # UUID-путь, не оригинальное имя
    return FileResponse(
        path=str(path),
        filename=rf.original_name,  # браузер видит оригинальное имя при сохранении
        media_type=rf.content_type,
    )""", "routers/reports.py")

    # ══════════════════════════════════════════
    # 8. ЧЕКЛИСТ И КАТЕГОРИИ
    # ══════════════════════════════════════════
    doc.add_heading("8. Чеклист и категории", 1)

    doc.add_paragraph(
        "Для каждой дисциплины заведующий формирует структуру чеклиста: "
        "создаёт категории (например, «Рабочая программа»), внутри — пункты (текстовые задачи), "
        "отмечает выполненные галочкой и прикрепляет файлы прямо к категории."
    )

    doc.add_heading("8.1 Структура данных", 2)
    add_code_block(doc, """\
Discipline (дисциплина)
  └── ChecklistCategory (категория) [один-ко-многим]
        ├── ChecklistItem (пункт) [один-ко-многим]
        │     ├── title: str
        │     ├── is_done: bool
        │     └── comment: str | None
        └── ReportFile (файл, category_id = cat.id) [через FK]""")

    doc.add_heading("8.2 Отображение файлов в категории", 2)
    doc.add_paragraph(
        "Файлы привязываются к категории через nullable FK category_id. "
        "В шаблоне files категории подгружаются через selectinload:"
    )
    add_code_block(doc, """\
# services/reports.py
result = await db.execute(
    select(ChecklistCategory)
    .options(
        selectinload(ChecklistCategory.items),
        selectinload(ChecklistCategory.files),  # ← файлы внутри категории
    )
    .where(ChecklistCategory.discipline_id == discipline_id)
)

# templates/reports/checklist.html
{% for cat in categories %}
  <div class="card">
    <h3>{{ cat.name }}</h3>
    {% for item in cat.items %}...{% endfor %}
    {% for f in cat.files %}
      <a href="/reports/files/{{ f.id }}/download">{{ f.original_name }}</a>
    {% endfor %}
  </div>
{% endfor %}""", "сервис + шаблон")

    # ══════════════════════════════════════════
    # 9. КОНФИГУРАЦИЯ И .env
    # ══════════════════════════════════════════
    doc.add_heading("9. Конфигурация и переменные окружения", 1)

    doc.add_paragraph(
        "Все чувствительные данные (строка подключения к БД, секретный ключ JWT) "
        "хранятся в файле .env, который не входит в git-репозиторий. "
        "Pydantic-Settings читает его автоматически:"
    )
    add_code_block(doc, """\
# .env
DATABASE_URL=postgresql+asyncpg://postgres:password@localhost:5432/edu_system
SECRET_KEY=my-super-secret-key-32-chars-min
ACCESS_TOKEN_EXPIRE_MINUTES=480""", ".env")

    add_code_block(doc, """\
# config.py
class Settings(BaseSettings):
    model_config = SettingsConfigDict(env_file=".env")

    DATABASE_URL: str = "postgresql+asyncpg://..."  # значение по умолчанию
    SECRET_KEY: str = "change-me-in-production"
    ACCESS_TOKEN_EXPIRE_MINUTES: int = 480  # 8 часов

settings = Settings()  # один экземпляр на всё приложение""", "config.py")

    add_note(doc, "Переменные окружения имеют приоритет над .env. "
             "На prod-сервере достаточно выставить переменные окружения — файл .env не нужен.")

    # ══════════════════════════════════════════
    # 10. КАК ЗАПУСТИТЬ
    # ══════════════════════════════════════════
    doc.add_heading("10. Как запустить проект локально", 1)

    steps_run = [
        ("Установить зависимости", "pip install -r requirements.txt"),
        ("Создать файл .env", "Скопировать .env.example → .env, прописать DATABASE_URL"),
        ("Создать БД в PostgreSQL", "createdb edu_system"),
        ("Применить миграции", "alembic upgrade head"),
        ("Заполнить тестовыми данными", "python seed.py"),
        ("Запустить сервер", "uvicorn app.main:app --reload"),
        ("Открыть в браузере", "http://localhost:8000  (логин: head@vgtu.ru / password123)"),
    ]

    table3 = doc.add_table(rows=len(steps_run) + 1, cols=2)
    table3.style = "Table Grid"
    table3.rows[0].cells[0].text = "Шаг"
    table3.rows[0].cells[1].text = "Команда / действие"
    for cell in table3.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, (desc, cmd) in enumerate(steps_run, 1):
        table3.rows[i].cells[0].text = desc
        p = table3.rows[i].cells[1].paragraphs[0]
        run = p.add_run(cmd)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    doc.add_paragraph()

    # ══════════════════════════════════════════
    # 11. КЛЮЧЕВЫЕ РЕШЕНИЯ
    # ══════════════════════════════════════════
    doc.add_heading("11. Ключевые архитектурные решения и компромиссы", 1)

    decisions = [
        (
            "SSR вместо SPA",
            "Jinja2 рендерит HTML на сервере — браузер получает готовую страницу.",
            "Нет лишней сборки, нет state management. Для CRUD-системы с ~5 страницами — идеально.",
            "Сложную интерактивность (drag-and-drop, real-time) сделать труднее."
        ),
        (
            "Cookie вместо localStorage для JWT",
            "Токен лежит в httponly-cookie.",
            "JS вообще не видит токен → XSS-атака не сможет его украсть.",
            "CSRF нужно учитывать (используем samesite=lax)."
        ),
        (
            "UUID-имена файлов",
            "Загруженный файл хранится под именем вида 'a1b2c3...pdf'.",
            "Исключает path-traversal атаки и коллизии имён.",
            "Имя файла не читаемо для человека — нужна таблица для сопоставления."
        ),
        (
            "expire_on_commit=False",
            "ORM-объекты не инвалидируются после commit().",
            "Можно читать атрибуты объекта после сохранения без доп. запроса.",
            "Объект может быть слегка устаревшим если данные менялись снаружи сессии."
        ),
        (
            "Selectinload вместо lazy load",
            "Связанные объекты загружаются явным SELECT IN запросом.",
            "Нет N+1 проблемы: 1 запрос на disciplines + 1 на teachers, а не 1 на каждую строку.",
            "Чуть больше данных загружается, если связанные объекты не всегда нужны."
        ),
        (
            "Клиентский поиск вместо серверного",
            "Фильтрация программ по названию/коду работает в JS без запроса на сервер.",
            "Мгновенный отклик, нет лишних HTTP-запросов при каждой букве.",
            "Не работает, если программ станет тысячи — тогда нужен серверный поиск."
        ),
    ]

    for dname, what, pro, con in decisions:
        doc.add_heading(dname, 3)
        t = doc.add_table(rows=3, cols=2)
        t.style = "Table Grid"
        labels = ["Что", "Плюс", "Компромисс"]
        values = [what, pro, con]
        for i, (lbl, val) in enumerate(zip(labels, values)):
            t.rows[i].cells[0].text = lbl
            t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            t.rows[i].cells[1].text = val
        doc.add_paragraph()

    # ══════════════════════════════════════════
    # 12. БЕЗОПАСНОСТЬ
    # ══════════════════════════════════════════
    doc.add_heading("12. Меры безопасности", 1)

    security_items = [
        ("bcrypt для паролей", "Медленный хэш — перебор 1 млн паролей занимает часы при утечке БД."),
        ("httponly cookie", "JavaScript не может прочитать токен — защита от XSS."),
        ("samesite=lax", "Браузер не отправляет cookie в cross-site запросах — защита от CSRF."),
        ("JWT подпись HS256", "Токен подписан SECRET_KEY — нельзя подделать без знания ключа."),
        ("Проверка MIME + размера", "Файл 50 МБ или .exe не пройдёт валидацию."),
        ("UUID-имена файлов", "Оригинальное имя не используется в пути — нет path-traversal."),
        ("Двухуровневая RBAC", "Проверка роли и в DI (get_head), и в бизнес-логике (approve_report)."),
        ("401 → редирект на логин", "Незалогиненный пользователь не видит JSON-ошибку, а попадает на страницу входа."),
    ]

    table4 = doc.add_table(rows=len(security_items) + 1, cols=2)
    table4.style = "Table Grid"
    table4.rows[0].cells[0].text = "Мера"
    table4.rows[0].cells[1].text = "Что защищает"
    for cell in table4.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, (measure, desc) in enumerate(security_items, 1):
        table4.rows[i].cells[0].text = measure
        table4.rows[i].cells[1].text = desc

    doc.add_paragraph()

    # ══════════════════════════════════════════
    # КОНЕЦ
    # ══════════════════════════════════════════
    doc.add_paragraph()
    p = doc.add_paragraph("Документ сгенерирован автоматически на основе исходного кода проекта.")
    p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    p.runs[0].font.size = Pt(9)

    out_path = "/Users/tikhonenko/Desktop/Описание программы.docx"
    doc.save(out_path)
    print(f"Сохранено: {out_path}")


if __name__ == "__main__":
    main()
